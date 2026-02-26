"""Unified local file reader — supports txt, md, csv, pdf, doc, docx."""

from __future__ import annotations

from pathlib import Path

from social_plugin.db import Database
from social_plugin.sources.models import SourceDocument
from social_plugin.utils.logger import get_logger

logger = get_logger()

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".rst", ".text", ".csv"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".doc", ".docx"}
ALL_SUPPORTED = TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS


def _read_text(path: Path) -> str:
    """Read a plain text file."""
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _read_docx(path: Path) -> str:
    """Extract text from a .doc/.docx file."""
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


class LocalReader:
    """Read content from local files (txt, md, csv, pdf, doc, docx)."""

    def __init__(self, db: Database):
        self.db = db

    def read(self, file_path: str, name: str = "") -> SourceDocument:
        """Read a local file, auto-detecting type by extension."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Local file not found: {path}")

        suffix = path.suffix.lower()
        logger.info("Reading local file: %s (type: %s)", path, suffix)

        if suffix in TEXT_EXTENSIONS:
            content = _read_text(path)
        elif suffix in PDF_EXTENSIONS:
            content = _read_pdf(path)
        elif suffix in DOCX_EXTENSIONS:
            content = _read_docx(path)
        else:
            logger.warning("Unsupported file type %s for %s — supported: %s", suffix, path, ", ".join(sorted(ALL_SUPPORTED)))
            content = f"[Unsupported file type: {path.name}]"

        title = name or path.stem

        source = SourceDocument(
            source_type="local_file",
            source_path=str(path),
            title=title,
            content=content,
        )
        source.compute_hash()

        existing = self.db.get_source_document_by_path(str(path))
        if existing and existing["content_hash"] == source.content_hash:
            logger.info("Local file %s unchanged, skipping", title)
            return SourceDocument.from_db_row(existing)

        self.db.insert_source_document(source.to_db_dict())
        logger.info("Stored local file: %s (%d chars)", title, len(content))
        return source

    def _discover_files(self, folder: Path) -> list[Path]:
        """Discover all supported files in a folder (non-recursive)."""
        files = [f for f in folder.iterdir() if f.is_file() and f.suffix.lower() in ALL_SUPPORTED]
        files.sort(key=lambda f: f.name)
        return files

    def read_all(self, config_sources: list[dict | str]) -> list[SourceDocument]:
        """Read all local files/folders from config.

        Each entry can be:
          - {"path": "/path/to/file.pdf", "name": "..."}
          - {"path": "/path/to/folder/"}       (reads all supported files in folder)
          - "/path/to/file_or_folder"          (bare string)
        """
        docs: list[SourceDocument] = []
        for src in config_sources:
            if isinstance(src, str):
                file_path = src
                name = ""
            else:
                file_path = src.get("path", "")
                name = src.get("name", "")
            if not file_path:
                continue

            path = Path(file_path)
            if path.is_dir():
                # Discover all supported files in folder
                discovered = self._discover_files(path)
                logger.info("Discovered %d supported files in %s", len(discovered), path)
                for f in discovered:
                    try:
                        doc = self.read(str(f))
                        docs.append(doc)
                    except Exception as e:
                        logger.error("Failed to read %s: %s", f, e)
            else:
                try:
                    doc = self.read(file_path, name)
                    docs.append(doc)
                except Exception as e:
                    logger.error("Failed to read local file %s: %s", file_path, e)
        return docs
